import os
import re
from pathlib import Path
from typing import Dict, List, Optional

import requests
from PyPDF2 import PdfReader
from docx import Document
from docx.shared import Pt
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.units import inch



def extract_text_from_pdf(filepath: str) -> str:
    text_parts = []
    with open(filepath, 'rb') as fh:
        reader = PdfReader(fh)
        for page in reader.pages:
            text_parts.append(page.extract_text() or '')
    return '\n'.join(text_parts).strip()


def extract_text_from_docx(filepath: str) -> str:
    document = Document(filepath)
    return '\n'.join(paragraph.text for paragraph in document.paragraphs).strip()


def extract_resume_text(filepath: str) -> str:
    extension = Path(filepath).suffix.lower()
    if extension == '.pdf':
        return extract_text_from_pdf(filepath)
    if extension == '.docx':
        return extract_text_from_docx(filepath)
    raise ValueError('Unsupported resume format')


def _tokenize(text: str) -> set:
    return {token.lower() for token in re.findall(r"[a-zA-Z0-9+#.]+", text) if token}


def _match_score(resume_text: str, job_description: Optional[str]) -> Dict[str, object]:
    if not job_description:
        return {'match_score': 0, 'matched_skills': [], 'match_summary': 'No job description provided.'}

    resume_tokens = _tokenize(resume_text)
    job_tokens = _tokenize(job_description)
    overlap = sorted(resume_tokens.intersection(job_tokens))
    if not overlap:
        return {'match_score': 0, 'matched_skills': [], 'match_summary': 'No overlapping skills or keywords found.'}

    score = min(100, int((len(overlap) / max(1, len(job_tokens))) * 100))
    return {
        'match_score': score,
        'matched_skills': overlap[:12],
        'match_summary': f"Matched {len(overlap)} relevant keywords from the job description.",
    }


def analyze_resume_text(text: str, job_description: Optional[str] = None) -> Dict[str, object]:
    words = [token.strip('.,()[]').lower() for token in text.split() if token.strip('.,()[]')]
    unique = set(words)
    skill_count = sum(1 for word in unique if word in {'python', 'sql', 'flask', 'docker', 'react', 'aws'})
    match_info = _match_score(text, job_description)
    return {
        'word_count': len(words),
        'skill_count': skill_count,
        'summary': text[:320] + '...' if len(text) > 320 else text,
        'match_score': match_info['match_score'],
        'matched_skills': match_info['matched_skills'],
        'match_summary': match_info['match_summary'],
    }


def generate_improvement_suggestions(text: str, job_description: Optional[str]) -> List[str]:
    if not job_description:
        return ['Add a job description to receive targeted improvement suggestions.']

    resume_tokens = _tokenize(text)
    job_tokens = _tokenize(job_description)
    missing_keywords = sorted(job_tokens.difference(resume_tokens))[:8]
    suggestions = []
    if missing_keywords:
        suggestions.append('Add these keywords to your resume: ' + ', '.join(missing_keywords))
    suggestions.append('Strengthen your project descriptions with measurable impact and role-specific outcomes.')
    suggestions.append('Add a professional summary tailored to this target role.')
    suggestions.append('Highlight tools and frameworks from the job description such as Python, SQL, Flask, Docker, AWS, or React.')
    return suggestions


def build_updated_resume_text(text: str, job_description: Optional[str]) -> str:
    suggestions = generate_improvement_suggestions(text, job_description)
    updated_sections = [
        'Updated Resume Draft',
        '',
        'Professional Summary',
        'Results-driven professional with strong experience in software development, data workflows, and product delivery.',
        '',
        'Suggested Improvements',
    ]
    for idx, suggestion in enumerate(suggestions, start=1):
        updated_sections.append(f"{idx}. {suggestion}")
    updated_sections.extend(['', 'Original Resume Content', text[:2400]])
    return '\n'.join(updated_sections)


def export_updated_resume(filepath: str, text: str, job_description: Optional[str]) -> str:
    updated_text = build_updated_resume_text(text, job_description)
    with open(filepath, 'w', encoding='utf-8') as fh:
        fh.write(updated_text)
    return filepath


def export_updated_docx(filepath: str, resume_text: str) -> str:
    """Export plain text to DOCX (best-effort formatting)."""
    doc = Document()

    # Very lightweight formatting: headings if line looks like a header.
    for line in resume_text.splitlines():
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph('')
            continue
        if stripped.lower() in {
            'updated resume draft',
            'professional summary',
            'suggested improvements',
            'original resume content',
        }:
            doc.add_heading(stripped, level=1)
        elif re.match(r'^\d+\.\s+', stripped):
            # bullet-like suggestions
            p = doc.add_paragraph(stripped)
            p.style = 'List Bullet'
        else:
            doc.add_paragraph(stripped)

    os.makedirs(os.path.dirname(filepath) or '.', exist_ok=True)
    doc.save(filepath)
    return filepath


def export_updated_pdf(filepath: str, resume_text: str) -> str:
    """Export plain text to PDF using reportlab (word-wrapping)."""
    os.makedirs(os.path.dirname(filepath) or '.', exist_ok=True)

    c = canvas.Canvas(filepath, pagesize=letter)
    width, height = letter
    left = 0.75 * inch
    top = height - 0.75 * inch
    y = top

    # Basic font settings
    c.setFont("Helvetica", 10)

    def draw_wrapped(text: str, x: float, y_pos: float, max_width: float):
        words = text.split(' ')
        line = ''
        lines = []
        for w in words:
            test = (line + ' ' + w).strip()
            if c.stringWidth(test, "Helvetica", 10) <= max_width:
                line = test
            else:
                lines.append(line)
                line = w
        if line:
            lines.append(line)
        for ln in lines:
            yield ln

    max_width = width - 2 * left

    for raw_line in resume_text.splitlines():
        line = raw_line.strip('\r')
        if not line.strip():
            y -= 12
            continue

        for wrapped_line in draw_wrapped(line, left, y, max_width):
            if y < 0.75 * inch:
                c.showPage()
                c.setFont("Helvetica", 10)
                y = top
            c.drawString(left, y, wrapped_line)
            y -= 12

    c.save()
    return filepath

